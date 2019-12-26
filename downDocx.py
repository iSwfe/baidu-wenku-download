#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import urllib.parse as urllibParse
import urllib.request as urllibReq
import json
import sys
from importlib import reload
# word操作库python-docx
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import config
import log

reload(sys)
# sys.setdefaultencoding('utf-8')

# python-docx 格式参考：https://zhuanlan.zhihu.com/p/23708800?utm_source=tuicool&utm_medium=referral

logger = log.Log(config.log_dir, config.log_name)


class DownDocx(object):
	def __init__(self, fileDir, url, info):
		self.fileDir = './' + fileDir + '/'
		self.URL = url
		self.WkInfo = info
		if not os.path.exists(self.fileDir):
			os.mkdir(self.fileDir)

	# 获取每页URL
	@staticmethod
	def geturl(urls, index):
		for url in urls:
			if url['pageIndex'] == index:
				return url['pageLoadUrl']
		return ''

	@staticmethod
	def encode(url, key):
		keyBegin = url.find(key)
		valueBegin = url.find('=', keyBegin) + 1
		valueEnd = url.find('&', valueBegin)
		return url[:valueBegin] + urllibParse.quote(url[valueBegin:valueEnd]) + url[valueEnd:]

	def down(self, pageStart, pageEnd):
		reqHeader = config.reqHeaderBDWK
		reqHeader['Referer'] = self.URL
		pageStart = 1 if pageStart is None else pageStart
		docFileName = self.fileDir + self.WkInfo['title'] + '.' + self.WkInfo['docType']
		document = Document()  # 创建doc(x)文档
		while pageStart <= pageEnd:
			jsonUrl = DownDocx.geturl(self.WkInfo['htmlUrls']['json'], pageStart)
			if len(jsonUrl) == 0:
				logger.error('下载文档失败，查找URL失败！')
				document.save(docFileName)
				return False
			jsonUrl = jsonUrl.replace('\\/', '/')

			jsonUrl = DownDocx.encode(jsonUrl, 'responseCacheControl')
			jsonUrl = DownDocx.encode(jsonUrl, 'responseExpires')
			jsonUrl = DownDocx.encode(jsonUrl, 'authorization')
			jsonUrl = DownDocx.encode(jsonUrl, 'x-bce-range')
			jsonUrl = DownDocx.encode(jsonUrl, 'token')
			print('第' + str(pageStart) + '页\t' + jsonUrl)

			req = urllibReq.Request(jsonUrl, headers=reqHeader)
			res = urllibReq.urlopen(req)
			res = res.read().decode()
			jsonRet = res[res.find('(') + 1: res.rfind(')')]
			logger.info('打印一下，获取json数据内容为 ' + jsonRet)
			jsonRet = json.loads(jsonRet)
			# 再处理获取的页面内容
			first = True
			for item in jsonRet['body']:
				if item['t'] != 'word':
					continue
				newLine = 'ps' in item and item['ps'] is not None and '_enter' in item['ps']
				if first or newLine:
					first = False
					pg = document.add_paragraph()
				if newLine:
					continue
				run = pg.add_run(item['c'])
				# 添加格式；分析不出来，就统一宋体、五号
				run.font.name = u'宋体'
				run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
				run.font.size = Pt(item['p']['h'])
			# 下一页
			if pageStart < pageEnd:
				document.add_page_break()
			pageStart += 1
		document.save(docFileName)
		return True
